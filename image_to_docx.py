import os
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor

def extract_text_from_image(image_path):
    """Extract text from an image using OCR."""
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        print(f"Error processing {image_path}: {str(e)}")
        return ""

def create_formatted_document(output_path):
    """Create a new Word document with specified formatting."""
    doc = Document()
    
    # Set page size to Legal (8.5 × 14 Inches)
    section = doc.sections[0]
    section.page_height = Inches(14)
    section.page_width = Inches(8.5)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0  # Single line spacing
    
    return doc

def process_images_to_docx():
    """Process all images in the Images folder and create a formatted Word document."""
    # Create the document
    doc = create_formatted_document('output.docx')
    
    # Get all image files from the Images folder
    image_extensions = ('.png', '.jpg', '.jpeg', '.bmp', '.tiff')
    images_folder = 'Images'
    
    # Process each image
    for filename in os.listdir(images_folder):
        if filename.lower().endswith(image_extensions):
            image_path = os.path.join(images_folder, filename)
            
            # Extract text from image
            text = extract_text_from_image(image_path)
            
            if text:
                # Add filename as heading
                doc.add_heading(filename, level=1)
                
                # Add extracted text
                paragraph = doc.add_paragraph(text)
                paragraph.style = doc.styles['Normal']
    
    # Save the document
    doc.save('output.docx')
    print("Document has been created successfully as 'output.docx'")

if __name__ == "__main__":
    process_images_to_docx() 